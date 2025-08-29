import streamlit as st
import pandas as pd
import os
import pandas as pd
import os
import tempfile
import shutil
import subprocess
import sys
from pathlib import Path
import zipfile
import io
import traceback
from datetime import datetime

# Global switch: turn conversation chunking on/off
ENABLE_CHUNKING = True  # Set to True to create _chunked.csv and other chunk files

# No config import needed for the reverted approach
# Load environment variables from .env file (for local development)
# For Streamlit Cloud, secrets are handled via st.secrets
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # dotenv not available in Streamlit Cloud

# Initialize session state for logging
def init_logging():
    """Initialize logging containers in session state"""
    if 'process_logs' not in st.session_state:
        st.session_state.process_logs = []
    if 'error_logs' not in st.session_state:
        st.session_state.error_logs = []
    if 'log_container' not in st.session_state:
        st.session_state.log_container = None
    if 'tx_logs' not in st.session_state:
        st.session_state.tx_logs = ""
    if 'ck_logs' not in st.session_state:
        st.session_state.ck_logs = ""
    if 'gap_log_messages' not in st.session_state:
        st.session_state.gap_log_messages = []

def add_log(message, log_type="INFO"):
    """Add a log message to the session state"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] [{log_type}] {message}"
    
    # Add to appropriate log collection
    if not hasattr(st.session_state, 'process_logs'):
        st.session_state.process_logs = []
    if not hasattr(st.session_state, 'error_logs'):
        st.session_state.error_logs = []
        
    st.session_state.process_logs.append(log_entry)
    if log_type in ["ERROR", "WARNING"]:
        st.session_state.error_logs.append(log_entry)

def display_logs():
    """Display all logs from session state in the original format"""
    # Transcription Logs
    if hasattr(st.session_state, 'tx_logs') and st.session_state.tx_logs:
        with st.expander("🔍 Transcription Logs", expanded=True):
            # Parse and display the JSON response in a readable format
            try:
                # Try to extract the actual log content if it's in JSON format
                if isinstance(st.session_state.tx_logs, str) and st.session_state.tx_logs.strip().startswith('{'):
                    import json
                    log_data = json.loads(st.session_state.tx_logs)
                    if isinstance(log_data, dict):
                        for key, value in log_data.items():
                            st.write(f"{key}: {value}")
                else:
                    # Display as regular text if not JSON
                    st.text_area("Status", value=st.session_state.tx_logs, height=150, disabled=True, key="tx_logs_display")
            except:
                # Fallback to raw text display if JSON parsing fails
                st.text_area("Status", value=st.session_state.tx_logs, height=150, disabled=True, key="tx_logs_fallback")

    # Gap Detection Logs
    if hasattr(st.session_state, 'gap_log_messages') and st.session_state.gap_log_messages:
        with st.expander("⏱️ Gap Detection Logs", expanded=True):
            if isinstance(st.session_state.gap_log_messages, list):
                for message in st.session_state.gap_log_messages:
                    st.write(message)
            else:
                st.write(st.session_state.gap_log_messages)

    # Error Logs
    if hasattr(st.session_state, 'error_logs') and st.session_state.error_logs:
        with st.expander("⚠️ Error Logs", expanded=True):
            for error in st.session_state.error_logs:
                st.error(error)

    # Process Logs (if any additional logs exist)
    if hasattr(st.session_state, 'process_logs') and st.session_state.process_logs:
        with st.expander("📋 Process Logs", expanded=True):
            st.text_area(
                "Processing History",
                value="\n".join(st.session_state.process_logs),
                height=200,
                disabled=True,
                key="process_logs_display"
            )

def create_logs_excel(output_folder, base_name, process_logs, error_logs):
    """Create Excel file with separate tabs for process and error logs"""
    try:
        # Create DataFrames for logs
        process_df = pd.DataFrame({'Log_Entry': process_logs})
        error_df = pd.DataFrame({'Error_Log_Entry': error_logs})
        
        # Create Excel file with multiple sheets
        excel_path = os.path.join(output_folder, f"{base_name}_logs.xlsx")
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            process_df.to_excel(writer, sheet_name='Process_Logs', index=False)
            error_df.to_excel(writer, sheet_name='Error_Logs', index=False)
        
        return excel_path
    except Exception as e:
        add_log(f"Failed to create logs Excel: {str(e)}", "ERROR")
        return None

# Optional DOCX generation
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Define format_timestamp function locally to avoid circular import
def format_timestamp(seconds):
    """Format seconds to HH:MM:SS,mmm format"""
    from datetime import timedelta
    td = timedelta(seconds=float(seconds))
    total_seconds = int(td.total_seconds())
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds_ = total_seconds % 60
    milliseconds = int((td.microseconds / 1000))
    return f"{hours:02d}:{minutes:02d}:{seconds_:02d},{milliseconds:03d}"

def get_api_key():
    """Get API key from Streamlit secrets or environment variables"""
    # Try Streamlit secrets first (for cloud deployment)
    try:
        if hasattr(st, 'secrets') and 'ASSEMBLYAI_API_KEY' in st.secrets:
            return st.secrets['ASSEMBLYAI_API_KEY']
    except:
        pass
    
    # Fallback to environment variable (for local development)
    api_key = os.getenv("ASSEMBLYAI_API_KEY")
    if api_key:
        return api_key
    
    # No API key found
    return None

def run_elevenlabs_transcription(audio_file_path, output_folder):
    """Run the elevenlabscribe.py script to process audio"""
    try:
        add_log("Starting AssemblyAI transcription process")
        
        # Get the absolute path to the audioUI directory
        audioui_dir = os.path.dirname(os.path.abspath(__file__))
        script_path = os.path.join(audioui_dir, "assemblyscribe.py")
        
        add_log(f"Using script path: {script_path}")
        
        # Create a temporary modified version of assemblyscribe.py
        with open(script_path, 'r', encoding='utf-8') as f:
            script_content = f.read()
        
        # Replace the hardcoded paths with absolute paths
        # Build path strings
        abs_audio = os.path.abspath(audio_file_path)
        abs_out   = os.path.abspath(output_folder)

        add_log(f"Processing audio file: {abs_audio}")
        add_log(f"Output folder: {abs_out}")

        # Prepend explicit vars to ensure they exist even if originals are commented out
        # Use raw strings to avoid Unicode escape issues with Windows paths
        preface = (
            f'AUDIO_FILE = r"{abs_audio}"\n'
            f'OUTPUT_FOLDER = r"{abs_out}"\n'
        )

        # Remove any existing hard-coded definitions (commented or not) to avoid confusion
        cleaned = []
        for line in script_content.splitlines():
            # Remove lines that assign AUDIO_FILE or OUTPUT_FOLDER
            if line.strip().startswith('AUDIO_FILE') or line.strip().startswith('OUTPUT_FOLDER'):
                continue
            cleaned.append(line)
        modified_content = preface + '\n'.join(cleaned)
        
        # Write temporary script with UTF-8 encoding
        temp_script = os.path.join(output_folder, "temp_assemblyscribe.py")
        with open(temp_script, 'w', encoding='utf-8') as f:
            f.write(modified_content)
        
        add_log("Created temporary script, executing transcription...")
        
        # Create environment for subprocess with API key
        env = os.environ.copy()
        api_key = get_api_key()
        if api_key:
            env["ASSEMBLYAI_API_KEY"] = api_key
        
        # Run the script from the audioUI directory
        result = subprocess.run([sys.executable, temp_script], 
                              capture_output=True, text=True, cwd=audioui_dir, env=env)
        
        # Clean up temp script
        os.remove(temp_script)
        
        log_text = f"STDOUT:\n{result.stdout}\n\nSTDERR:\n{result.stderr}"
        
        if result.returncode == 0:
            add_log("Transcription completed successfully!")
            st.success("Transcription completed successfully!")
            
            # Add quality checking after successful transcription
            try:
                add_log("Starting quality analysis...")
                base_name = os.path.splitext(os.path.basename(audio_file_path))[0]
                
                # Find the merged CSV file created by assemblyscribe.py
                merged_csv_path = os.path.join(output_folder, f"{base_name}_merged.csv")
                
                if os.path.exists(merged_csv_path):
                    add_log(f"Found merged CSV for quality analysis: {merged_csv_path}")
                    
                    # Apply quality checks
                    add_log("Normalizing speaker names...")
                    normalize_speaker_names(merged_csv_path)
                    
                    add_log("Validating language and marking segments...")
                    validate_language_and_mark(merged_csv_path)
                    
                    add_log("Detecting gaps in transcript...")
                    gap_count = flag_long_gaps(merged_csv_path, gap_seconds=60)
                    
                    # Run quality analysis
                    df = pd.read_csv(merged_csv_path)
                    label_col = 'Label' if 'Label' in df.columns else ('Notes' if 'Notes' in df.columns else None)
                    
                    # Store quality analysis results in session state
                    quality_results = {
                        'flagged_segments': [],
                        'long_segments': [],
                        'gap_segments': [],
                        'total_issues': 0,
                        'total_segments': len(df)
                    }
                    
                    if label_col:
                        # Check for gibberish/unsupported language - fix .str accessor error
                        try:
                            flagged = df[df[label_col].astype(str).str.contains('require human transcription', na=False)]
                            if not flagged.empty:
                                quality_results['flagged_segments'] = flagged.to_dict('records')
                        except Exception as e:
                            add_log(f"Warning: Could not check flagged segments: {e}")
                        
                        # Check for long segments (>120s) - fix .str accessor error
                        try:
                            long_segments = df[df[label_col].astype(str).str.contains('>120s single segment', na=False)]
                            if not long_segments.empty:
                                quality_results['long_segments'] = long_segments.to_dict('records')
                        except Exception as e:
                            add_log(f"Warning: Could not check long segments: {e}")
                        
                        # Check for gaps - fix .str accessor error
                        try:
                            gaps = df[df[label_col].astype(str).str.contains('>60s silence', na=False)]
                            if gap_count > 0:
                                quality_results['gap_segments'] = gaps.to_dict('records')
                        except Exception as e:
                            add_log(f"Warning: Could not check gap segments: {e}")
                        
                        # Summary - fix variable reference errors
                        try:
                            total_issues = len(quality_results.get('flagged_segments', [])) + len(quality_results.get('long_segments', [])) + gap_count
                            quality_results['total_issues'] = total_issues
                        except Exception as e:
                            quality_results['total_issues'] = 0
                            add_log(f"Warning: Could not calculate total issues: {e}")
                    
                    # Store quality results in session state
                    st.session_state['quality_results'] = quality_results
                    st.session_state['quality_df'] = df.to_dict('records')
                    st.session_state['label_col'] = label_col
                    
                    # Create Excel with quality analysis
                    excel_with_quality = create_transcript_with_quality_excel(
                        merged_csv_path, quality_results, label_col, output_folder, base_name
                    )
                    
                    # Check if chunked files exist and add quality analysis
                    chunked_csv_path = os.path.join(output_folder, f"{base_name}_chunked.csv")
                    if os.path.exists(chunked_csv_path):
                        add_log("Found chunked CSV, adding quality analysis...")
                        chunked_df = pd.read_csv(chunked_csv_path)
                        
                        # Run chunked quality analysis
                        chunked_quality_results = analyze_chunked_quality(chunked_df)
                        
                        # Debug: Print quality results
                        add_log(f"DEBUG: Quality analysis found {len(chunked_quality_results.get('long_gaps', []))} long gaps")
                        add_log(f"DEBUG: Quality analysis found {len(chunked_quality_results.get('short_chunks', []))} short chunks")
                        add_log(f"DEBUG: Quality analysis found {len(chunked_quality_results.get('missing_speakers', []))} missing speakers")
                        
                        # Ensure Notes column exists first
                        if 'Notes' not in chunked_df.columns:
                            chunked_df['Notes'] = ''
                            add_log("Added Notes column to chunked DataFrame")
                        
                        # Apply quality flags to chunked DataFrame
                        try:
                            chunked_df = apply_quality_flags_to_chunked_df(chunked_df, chunked_quality_results)
                            add_log("Applied quality flags to chunked data")
                        except Exception as e:
                            add_log(f"Error applying quality flags: {str(e)}", "ERROR")
                            add_log(f"Traceback: {traceback.format_exc()}", "ERROR")
                        
                                        # Quality flags are now applied by apply_quality_flags_to_chunked_df function
                        # No sample flags needed - real quality analysis is working
                        
                        # Create chunked Excel with quality analysis
                        chunked_excel_with_quality = os.path.join(output_folder, f"{base_name}_chunked_with_quality.xlsx")
                        create_chunked_excel_with_quality(
                            chunked_df, chunked_quality_results, chunked_excel_with_quality, 
                            audio_file_path, len(df)
                        )
                        add_log(f"Created chunked Excel with quality analysis: {chunked_excel_with_quality}")
                        
                        # Also save the updated chunked CSV with quality flags
                        chunked_df.to_csv(chunked_csv_path, index=False)
                        add_log(f"Updated chunked CSV with quality flags: {chunked_csv_path}")
                        
                        # Verify Notes column has data
                        notes_with_data = chunked_df[chunked_df['Notes'].notna() & (chunked_df['Notes'] != '')]
                        add_log(f"Final check: {len(notes_with_data)} chunks have quality flags in Notes column")
                        
                        # Debug: Show first few Notes entries
                        for i in range(min(3, len(chunked_df))):
                            note_content = chunked_df.iloc[i]['Notes']
                            add_log(f"Chunk {i+1} Notes: '{note_content}'")
                    
                    add_log("Quality analysis completed successfully!")
                    st.success("Quality analysis completed!")
                    
                else:
                    add_log("Merged CSV not found, skipping quality analysis", "WARNING")
                    
            except Exception as e:
                add_log(f"Quality analysis failed: {str(e)}", "ERROR")
                add_log(f"Quality analysis traceback: {traceback.format_exc()}", "ERROR")
                # Don't fail the entire transcription if quality analysis fails
                st.warning(f"Quality analysis failed: {str(e)}")
            
            return True, log_text
        else:
            add_log(f"Transcription failed with error code {result.returncode}", "ERROR")
            add_log(f"Error details: {log_text}", "ERROR")
            st.error("Transcription failed. Check logs for details.")
            return False, log_text
            
    except Exception as e:
        error_msg = f"Error running transcription: {str(e)}"
        add_log(error_msg, "ERROR")
        add_log(f"Traceback: {traceback.format_exc()}", "ERROR")
        st.error(error_msg)
        st.error(f"Traceback: {traceback.format_exc()}")
        return False, f"Exception: {str(e)}"

def run_chunking(csv_file_path, output_folder):
    """Run the chunk.py script to create conversation chunks"""
    try:
        add_log("Starting chunking process")
        
        # Get the absolute path to the audioUI directory
        audioui_dir = os.path.dirname(os.path.abspath(__file__))
        script_path = os.path.join(audioui_dir, "chunk.py")
        
        add_log(f"Using chunking script: {script_path}")
        add_log(f"Input CSV: {csv_file_path}")
        
        # Create a temporary modified version of chunk.py
        with open(script_path, 'r') as f:
            script_content = f.read()
        
        # Replace the hardcoded paths with absolute paths
        modified_content = script_content.replace(
            'INPUT_CSV = "/home/alexong/intage/1476_JunKaiOng_GEFA_m2_merged.csv"',
            f'INPUT_CSV = "{os.path.abspath(csv_file_path)}"'
        ).replace(
            'OUTPUT_FOLDER = "/home/alexong/intage"',
            f'OUTPUT_FOLDER = "{os.path.abspath(output_folder)}"'
        )
        
        # Also replace any other hardcoded paths that might exist
        modified_content = modified_content.replace(
            '/home/alexong/intage',
            f'{os.path.abspath(output_folder)}'
        )
        
        # Write temporary script
        temp_script = os.path.join(output_folder, "temp_chunk.py")
        with open(temp_script, 'w') as f:
            f.write(modified_content)
        
        add_log("Created temporary chunking script, executing...")
        
        # Run the script from the audioUI directory
        result = subprocess.run([sys.executable, temp_script], 
                              capture_output=True, text=True, cwd=audioui_dir)
        
        # Clean up temp script
        os.remove(temp_script)
        
        log_text = f"STDOUT:\n{result.stdout}\n\nSTDERR:\n{result.stderr}"
        if result.returncode == 0:
            add_log("Chunking completed successfully!")
            st.success("Chunking completed successfully!")
            return True, log_text
        else:
            add_log("Chunking failed", "ERROR")
            add_log(f"Chunking error details: {log_text}", "ERROR")
            st.error("Chunking failed")
            return False, log_text
            
    except Exception as e:
        error_msg = f"Error running chunking: {str(e)}"
        add_log(error_msg, "ERROR")
        add_log(f"Chunking traceback: {traceback.format_exc()}", "ERROR")
        st.error(error_msg)
        st.error(f"Traceback: {traceback.format_exc()}")
        return False, f"Exception: {str(e)}"

# Enhanced find_output_files: accept both _chunked and _merged_chunked naming
def find_output_files(output_folder, base_name):
    """Find the generated output files"""
    # The assemblyscribe.py creates files like: base_name_merged.csv
    # The chunk.py creates files like: base_name_merged_chunked.csv
    files = {
        'merged_csv': os.path.join(output_folder, f"{base_name}_merged.csv"),
        'chunked_csv': os.path.join(output_folder, f"{base_name}_merged_chunked.csv"),
        'chunked_excel': os.path.join(output_folder, f"{base_name}_merged_chunked.xlsx"),
        'chunked_txt': os.path.join(output_folder, f"{base_name}_merged_chunked.txt")
    }
    # Fallback names
    fallback = {
        # same naming without "merged_" prefix
        'chunked_csv': os.path.join(output_folder, f"{base_name}_chunked.csv"),
        'chunked_excel': os.path.join(output_folder, f"{base_name}_chunked.xlsx"),
        'chunked_txt': os.path.join(output_folder, f"{base_name}_chunked.txt"),
        # full-file path variants (previous implementation)
        'chunked_csv_full': os.path.join(output_folder, f"{base_name}_full_chunked.csv"),
        'chunked_excel_full': os.path.join(output_folder, f"{base_name}_full_chunked.xlsx"),
        'chunked_txt_full': os.path.join(output_folder, f"{base_name}_full_chunked.txt"),
        # NEW: full-file *merged* variants
        'merged_csv_full': os.path.join(output_folder, f"{base_name}_full_merged.csv"),
        'chunked_csv_full_merged': os.path.join(output_folder, f"{base_name}_full_merged_chunked.csv"),
        'chunked_excel_full_merged': os.path.join(output_folder, f"{base_name}_full_merged_chunked.xlsx"),
        'chunked_txt_full_merged': os.path.join(output_folder, f"{base_name}_full_merged_chunked.txt")
    }
    
    existing_files = {}
    for file_type, file_path in files.items():
        if os.path.exists(file_path):
            existing_files[file_type] = file_path

    # If missing expected files, check fallback names including full variants
    for key, path in fallback.items():
        if os.path.exists(path):
            if 'merged_csv' in key:
                existing_files['merged_csv'] = path
            elif 'csv' in key:
                existing_files['chunked_csv'] = path
            elif 'excel' in key:
                existing_files['chunked_excel'] = path
            elif 'txt' in key:
                existing_files['chunked_txt'] = path

    if not existing_files:
        print(f"No expected files found. Files in {output_folder}:")
        for file in os.listdir(output_folder):
            print(f"  - {file}")
    
    return existing_files

# -------------------------------------------------------------
# Helper: ensure first speaker is Mysteryshopper, second is InsuranceAgent
# -------------------------------------------------------------

def normalize_speaker_names(csv_path):
    """Simple speaker mapping: alternating between Mysteryshopper and InsuranceAgent.
    Row 0: Mysteryshopper, Row 1: InsuranceAgent, Row 2: Mysteryshopper, etc."""
    try:
        df = pd.read_csv(csv_path)
        
        # Simple alternating pattern: even rows = Mysteryshopper, odd rows = InsuranceAgent
        df['Speaker'] = df.index.map(lambda i: 'Mysteryshopper' if i % 2 == 0 else 'InsuranceAgent')
        
        df.to_csv(csv_path, index=False)
        print(f"Applied simple alternating speaker mapping: {len(df)} segments")
        return True
    except Exception as e:
        print(f"Speaker normalization failed for {csv_path}: {e}")
        return False

# -------------------------------------------------------------
# Helper: validate language and flag segments needing manual review
# -------------------------------------------------------------
import re

def validate_language_and_mark(csv_path):
    """Detect whether each Text entry is English or Chinese. If not, mark Label/Notes column."""
    try:
        df = pd.read_csv(csv_path)
        # Ensure a column exists to hold the flag
        label_col = 'Label' if 'Label' in df.columns else ('Notes' if 'Notes' in df.columns else None)
        if label_col is None:
            # Add a new column if missing
            df['Label'] = ''
            label_col = 'Label'
        def is_en_or_zh(text: str) -> bool:
            text = str(text).strip()
            
            # Skip empty or very short text
            if len(text) < 2:
                return True  # Don't flag very short utterances
            
            # Flag replacement char counts (corrupted text)
            if text.count('\uFFFD') / max(1, len(text)) > 0.05:
                return False
            
            # Accept Chinese characters
            if re.search(r'[\u4e00-\u9fff]', text):
                return True
            
            # Check for common English words (case insensitive)
            common_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'among', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'what', 'where', 'when', 'why', 'how', 'who', 'which', 'that', 'this', 'these', 'those', 'a', 'an', 'my', 'your', 'his', 'her', 'its', 'our', 'their', 'me', 'you', 'him', 'her', 'us', 'them', 'i', 'he', 'she', 'it', 'we', 'they', 'yes', 'no', 'ok', 'okay', 'hello', 'hi', 'bye', 'thanks', 'thank', 'please', 'sorry', 'excuse', 'insurance', 'number', 'name', 'call', 'phone', 'contact', 'give', 'get', 'take', 'make', 'go', 'come', 'see', 'know', 'think', 'say', 'tell', 'ask', 'need', 'want', 'like', 'time', 'day', 'year', 'way', 'back', 'good', 'new', 'first', 'last', 'long', 'great', 'little', 'own', 'other', 'old', 'right', 'big', 'high', 'different', 'small', 'large', 'next', 'early', 'young', 'important', 'few', 'public', 'bad', 'same', 'able']
            text_lower = text.lower()
            for word in common_words:
                if word in text_lower:
                    return True
            
            # Remove punctuation and whitespace for character analysis
            _clean = re.sub(r'[^\w]', '', text)
            if not _clean:
                return True  # Only punctuation/whitespace, don't flag
            
            # Count different character types
            letters = len(re.findall(r'[A-Za-z]', _clean))
            digits = len(re.findall(r'\d', _clean))
            total_chars = len(_clean)
            
            # Accept if has reasonable amount of English letters (lowered threshold)
            if letters >= total_chars * 0.25:  # Reduced from 0.4 to 0.25
                return True
            
            # Accept if mostly numeric (phone numbers, account numbers, etc.)
            if digits >= total_chars * 0.5:  # Reduced from 0.6 to 0.5
                return True
            
            # Accept mixed alphanumeric content
            if (letters + digits) >= total_chars * 0.7:
                return True
            
            # If text contains any English letters and is reasonably short, accept it
            if letters > 0 and total_chars <= 20:
                return True
            
            return False
        df[label_col] = df['Text'].apply(lambda t: '' if is_en_or_zh(t) else 'require human transcription')
        df.to_csv(csv_path, index=False)
        return True
    except Exception as e:
        print(f"Language validation failed for {csv_path}: {e}")
        return False

# -------------------------------------------------------------
# Helper: create DOCX from chunked CSV
# -------------------------------------------------------------

def generate_docx_from_chunked_csv(csv_path, docx_path):
    """Create a Word document from the chunked CSV (if python-docx is available)."""
    if not DOCX_AVAILABLE:
        return False, "python-docx not installed"
    try:
        import docx
        df = pd.read_csv(csv_path)
        doc = docx.Document()
        doc.add_heading("Conversation Chunks", level=1)

        for idx, row in df.iterrows():
            doc.add_heading(f"Chunk {idx + 1}", level=2)
            doc.add_paragraph(f"Time Range: {row['Start Time']} - {row['End Time']}")
            combined_text = row['Combined Text'].replace('  \n', '\n')
            doc.add_paragraph(combined_text)
            doc.add_page_break()

        doc.save(docx_path)
        return True, "DOCX created"
    except Exception as e:
        return False, str(e)

def create_conversation_chunks(df):
    """
    Creates conversation chunks from a DataFrame.
    Each chunk contains a fixed number of segments (CHUNK_SIZE).
    The combined text includes speaker names.
    """
    CHUNK_SIZE = 5  # Fixed chunk size as requested
    
    # Convert to list of dictionaries for easier processing
    segments = df.to_dict('records')
    
    # Create chunks based on number of segments
    chunks = []
    current_chunk = []
    
    for segment in segments:
        current_chunk.append(segment)
        
        # If we've reached the desired number of segments, start a new chunk
        if len(current_chunk) >= CHUNK_SIZE:
            chunks.append(current_chunk)
            current_chunk = []
    
    # Add the last chunk if it has content
    if current_chunk:
        chunks.append(current_chunk)
    
    # Create the formatted output data (no Speakers/Label columns)
    formatted_chunks = []
    
    for chunk_idx, chunk in enumerate(chunks, 1):
        # Get chunk start and end times
        start_time = chunk[0]['Start Time']
        end_time = chunk[-1]['End Time']
        
        # Create combined text that already includes speaker names
        combined_text_parts = []
        for segment in chunk:
            speaker = segment['Speaker']
            text = segment['Text']
            combined_text_parts.append(f"{speaker}: {text}")
        
        combined_text = "\n".join(combined_text_parts)
        
        # Build final row (matching the successful chunked format)
        chunk_record = {
            'Chunk_Number': chunk_idx,
            'Start_Time': start_time,
            'End_Time': end_time,
            'Combined_Text': combined_text
        }
        
        formatted_chunks.append(chunk_record)
    
    return pd.DataFrame(formatted_chunks)

def analyze_chunked_quality(chunked_df):
    """
    Analyze quality issues in chunked transcript data.
    """
    print(f"DEBUG: Analyzing quality for {len(chunked_df)} chunks")
    
    quality_results = {
        'total_chunks': len(chunked_df),
        'long_gaps': [],
        'short_chunks': [],
        'missing_speakers': [],
        'total_issues': 0
    }
    
    for idx, row in chunked_df.iterrows():
        chunk_num = row['Chunk_Number']
        combined_text = str(row['Combined_Text'])
        start_time = row['Start_Time']
        end_time = row['End_Time']
        
        print(f"DEBUG: Analyzing chunk {chunk_num}: {combined_text[:50]}...")
        
        # Check for missing speaker labels (more flexible detection)
        speaker_patterns = ['Mysteryshopper:', 'InsuranceAgent:', 'Speaker A:', 'Speaker B:', 'SPEAKER_']
        has_speaker = any(pattern in combined_text for pattern in speaker_patterns)
        
        if not has_speaker:
            quality_results['missing_speakers'].append({
                'Chunk_Number': chunk_num,
                'Issue': 'No speaker labels found',
                'Start_Time': start_time,
                'End_Time': end_time
            })
            print(f"DEBUG: Found missing speaker in chunk {chunk_num}")
        
        # Check for very short chunks (less than 50 characters or 2 lines)
        text_lines = [line.strip() for line in combined_text.split('\n') if line.strip()]
        if len(combined_text.strip()) < 50 or len(text_lines) < 2:
            quality_results['short_chunks'].append({
                'Chunk_Number': chunk_num,
                'Issue': f'Short chunk ({len(combined_text)} chars, {len(text_lines)} lines)',
                'Start_Time': start_time,
                'End_Time': end_time
            })
            print(f"DEBUG: Found short chunk {chunk_num}: {len(combined_text)} chars")
        
        # Check for long gaps between chunks (>60s for more sensitivity)
        if idx > 0:
            prev_end = chunked_df.iloc[idx-1]['End_Time']
            # Convert time strings to seconds for comparison
            try:
                prev_end_sec = time_to_seconds(prev_end)
                start_sec = time_to_seconds(start_time)
                gap = start_sec - prev_end_sec
                
                print(f"DEBUG: Gap between chunks {chunk_num-1} and {chunk_num}: {gap:.1f}s")
                
                if gap > 240:  # 240 seconds = 4 minutes
                    quality_results['long_gaps'].append({
                        'Between_Chunks': f"{chunk_num-1} and {chunk_num}",
                        'Gap_Duration': f"{gap:.1f}s",
                        'Issue': f'Long gap of {gap:.1f}s between chunks',
                        'Previous_End': prev_end,
                        'Current_Start': start_time
                    })
                    print(f"DEBUG: Found long gap: {gap:.1f}s between chunks {chunk_num-1} and {chunk_num}")
            except Exception as e:
                print(f"Error calculating gap: {e}")
    
    # Calculate total issues
    quality_results['total_issues'] = (
        len(quality_results['long_gaps']) + 
        len(quality_results['short_chunks']) + 
        len(quality_results['missing_speakers'])
    )
    
    print(f"DEBUG: Quality analysis complete. Found {quality_results['total_issues']} total issues:")
    print(f"  - {len(quality_results['long_gaps'])} long gaps")
    print(f"  - {len(quality_results['short_chunks'])} short chunks")
    print(f"  - {len(quality_results['missing_speakers'])} missing speakers")
    
    return quality_results

def time_to_seconds(time_str):
    """Convert time string (HH:MM:SS,mmm) to seconds"""
    try:
        if ',' in time_str:
            time_part, ms_part = time_str.split(',')
            ms = int(ms_part) / 1000
        else:
            time_part = time_str
            ms = 0
        
        parts = time_part.split(':')
        hours = int(parts[0])
        minutes = int(parts[1])
        seconds = int(parts[2])
        
        return hours * 3600 + minutes * 60 + seconds + ms
    except:
        return 0

def apply_quality_flags_to_chunked_df(chunked_df, quality_results):
    """
    Apply quality analysis flags directly to the chunked DataFrame.
    """
    print(f"DEBUG: Applying quality flags. Quality results: {quality_results}")
    
    # Ensure Notes column exists and is properly initialized
    if 'Notes' not in chunked_df.columns:
        chunked_df['Notes'] = ''
        print("DEBUG: Added Notes column")
    else:
        # Fill any NaN values in existing Notes column
        chunked_df['Notes'] = chunked_df['Notes'].fillna('')
        print("DEBUG: Filled NaN values in existing Notes column")
    
    flags_applied = 0
    
    # Flag long gaps - fix the logic
    for gap in quality_results.get('long_gaps', []):
        between_chunks = gap.get('Between_Chunks', '')
        print(f"DEBUG: Processing gap: {between_chunks}")
        
        # Parse "1 and 2" format
        import re
        match = re.search(r'(\d+) and (\d+)', between_chunks)
        if match:
            chunk1, chunk2 = int(match.group(1)), int(match.group(2))
            # Flag the second chunk (where the gap occurs before it)
            if chunk2 <= len(chunked_df):
                idx = chunk2 - 1  # Convert to 0-based index
                current_notes = chunked_df.iloc[idx]['Notes']
                flag_text = f"Long gap before this chunk ({gap.get('Gap_Duration', 'Unknown')})"
                
                if pd.isna(current_notes) or str(current_notes).strip() == '':
                    chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = flag_text
                else:
                    chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = f"{str(current_notes).strip()}; {flag_text}"
                
                flags_applied += 1
                print(f"DEBUG: Applied long gap flag to chunk {chunk2}")
    
    # Flag short chunks
    for short_chunk in quality_results.get('short_chunks', []):
        chunk_num = short_chunk.get('Chunk_Number', 0)
        print(f"DEBUG: Processing short chunk: {chunk_num}")
        
        if chunk_num > 0 and chunk_num <= len(chunked_df):
            idx = chunk_num - 1
            current_notes = chunked_df.iloc[idx]['Notes']
            flag_text = "Short chunk - may indicate transcription gaps"
            
            if pd.isna(current_notes) or str(current_notes).strip() == '' or current_notes == 'nan':
                chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = flag_text
            else:
                chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = f"{str(current_notes).strip()}; {flag_text}"
            
            flags_applied += 1
            print(f"DEBUG: Applied short chunk flag to chunk {chunk_num}")
    
    # Flag missing speakers
    for missing_speaker in quality_results.get('missing_speakers', []):
        chunk_num = missing_speaker.get('Chunk_Number', 0)
        print(f"DEBUG: Processing missing speaker: {chunk_num}")
        
        if chunk_num > 0 and chunk_num <= len(chunked_df):
            idx = chunk_num - 1
            current_notes = chunked_df.iloc[idx]['Notes']
            flag_text = "Missing speaker labels detected"
            
            if pd.isna(current_notes) or str(current_notes).strip() == '' or current_notes == 'nan':
                chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = flag_text
            else:
                chunked_df.iloc[idx, chunked_df.columns.get_loc('Notes')] = f"{str(current_notes).strip()}; {flag_text}"
            
            flags_applied += 1
            print(f"DEBUG: Applied missing speaker flag to chunk {chunk_num}")
    
    print(f"DEBUG: Total flags applied: {flags_applied}")
    return chunked_df

def create_chunked_excel_with_quality(chunked_df, quality_results, output_path, source_file, total_segments):
    """
    Create Excel file with chunked transcript and quality analysis sheets.
    """
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Sheet 1: Chunked Transcript
        chunked_df.to_excel(writer, sheet_name='Chunked_Transcript', index=False)
        
        # Sheet 2: Quality Summary
        quality_summary = []
        quality_summary.append(['Metric', 'Value'])
        quality_summary.append(['Source File', str(source_file)])
        quality_summary.append(['Total Original Segments', str(total_segments)])
        quality_summary.append(['Total Chunks Created', str(quality_results['total_chunks'])])
        quality_summary.append(['Segments per Chunk', str(5)])
        quality_summary.append(['Total Quality Issues', str(quality_results['total_issues'])])
        quality_summary.append(['Long Gaps (>240s)', str(len(quality_results['long_gaps']))])
        quality_summary.append(['Short Chunks', str(len(quality_results['short_chunks']))])
        quality_summary.append(['Missing Speaker Labels', str(len(quality_results['missing_speakers']))])
        
        quality_df = pd.DataFrame(quality_summary[1:], columns=quality_summary[0])
        quality_df.to_excel(writer, sheet_name='Quality_Summary', index=False)
        
        # Sheet 3: Quality Issues (if any)
        if quality_results['total_issues'] > 0:
            all_issues = []
            
            # Add long gaps
            for issue in quality_results['long_gaps']:
                all_issues.append({
                    'Issue_Type': 'Long Gap',
                    'Location': issue['Between_Chunks'],
                    'Description': issue['Issue'],
                    'Duration/Details': issue['Gap_Duration'],
                    'Start_Time': issue.get('Previous_End', ''),
                    'End_Time': issue.get('Current_Start', '')
                })
            
            # Add short chunks
            for issue in quality_results['short_chunks']:
                all_issues.append({
                    'Issue_Type': 'Short Chunk',
                    'Location': f"Chunk {issue['Chunk_Number']}",
                    'Description': issue['Issue'],
                    'Duration/Details': '',
                    'Start_Time': issue['Start_Time'],
                    'End_Time': issue['End_Time']
                })
            
            # Add missing speakers
            for issue in quality_results['missing_speakers']:
                all_issues.append({
                    'Issue_Type': 'Missing Speakers',
                    'Location': f"Chunk {issue['Chunk_Number']}",
                    'Description': issue['Issue'],
                    'Duration/Details': '',
                    'Start_Time': issue['Start_Time'],
                    'End_Time': issue['End_Time']
                })
            
            if all_issues:
                issues_df = pd.DataFrame(all_issues)
                issues_df.to_excel(writer, sheet_name='Quality_Issues', index=False)

def run_fullfile_transcription(audio_file, output_dir):
    """
    One-shot call to AssemblyAI for the whole file.
    Produces {base}_full.csv  (+ merged CSV) in output_dir.
    """
    import assemblyai as aai

    # Get API key from secrets or environment
    api_key = get_api_key()
    if not api_key:
        return False, "No API key available. Set ASSEMBLYAI_API_KEY in Streamlit secrets or .env file."
    
    try:
        debug_msg = f"Starting transcription for file: {audio_file}"
        add_log(debug_msg, "DEBUG")
        
        debug_msg = f"API key (first 8 chars): {api_key[:8]}..."
        add_log(debug_msg, "DEBUG")
        
        # Set API key
        aai.settings.api_key = api_key
        add_log("API key set successfully", "DEBUG")
        
        # Configure transcription settings
        config = aai.TranscriptionConfig(
            speaker_labels=True,  # Enable speaker diarization
            language_code="en",   # English language
            punctuate=True,       # Add punctuation
            format_text=True,     # Format text
        )
        add_log("Configuration created successfully", "DEBUG")
        
        transcriber = aai.Transcriber()
        add_log("Transcriber created successfully", "DEBUG")
        
        # Upload and transcribe
        add_log("Starting transcription request...", "DEBUG")
        transcript = transcriber.transcribe(audio_file, config)
        add_log(f"Transcription completed with status: {transcript.status}", "DEBUG")
        
        if transcript.status == aai.TranscriptStatus.error:
            error_msg = f"Transcription failed: {transcript.error}"
            add_log(error_msg, "ERROR")
            return False, error_msg
        
        log = f"Status: {transcript.status}\nTranscript ID: {transcript.id}"
        add_log(f"Success log: {log}", "DEBUG")
        
        segments = []
        
        # Process utterances (speaker-segmented results)
        if transcript.utterances:
            for utterance in transcript.utterances:
                segments.append({
                    "speaker_id": utterance.speaker,
                    "start": utterance.start / 1000.0,  # Convert ms to seconds
                    "end": utterance.end / 1000.0,
                    "text": utterance.text.strip()
                })
        
        elif transcript.words:
            # Fallback to word-level processing
            cur = None
            for word in transcript.words:
                speaker_id = getattr(word, 'speaker', 'A')
                if cur is None or cur["speaker_id"] != speaker_id:
                    if cur:
                        segments.append(cur)
                    cur = {"speaker_id": speaker_id, "start": word.start / 1000.0,
                           "end": word.end / 1000.0, "text": word.text}
                else:
                    cur["end"] = word.end / 1000.0
                    cur["text"] += " " + word.text
            if cur:
                segments.append(cur)
        
        if not segments:
            return False, "No segments found in transcription"

        # → DataFrame & CSV
        df = pd.DataFrame([{
            "Start Time": format_timestamp(s["start"]),
            "End Time":   format_timestamp(s["end"]),
            "Speaker":    f"SPEAKER_{s['speaker_id']}",
            "Text":       s["text"],
            "Notes":      ""
        } for s in segments])

        base = Path(audio_file).stem
        csv  = Path(output_dir) / f"{base}_full.csv"
        df.to_csv(csv, index=False)

        # Create merged version similar to assemblyscribe.py
        merged_rows = []
        prev_row = None
        for row in df.to_dict('records'):
            if prev_row is None:
                prev_row = row.copy()
            elif row['Speaker'] == prev_row['Speaker']:
                # Extend previous segment
                prev_row['End Time'] = row['End Time']
                prev_row['Text'] = f"{prev_row['Text']} {row['Text']}"
            else:
                merged_rows.append(prev_row)
                prev_row = row.copy()
        if prev_row is not None:
            merged_rows.append(prev_row)

        merged_df = pd.DataFrame(merged_rows)
        merged_csv = Path(output_dir) / f"{base}_full_merged.csv"
        merged_df.to_csv(merged_csv, index=False)
        
        # Add chunking integration
        if ENABLE_CHUNKING:
            try:
                add_log(f"CHUNKING ENABLED - Starting chunking process...", "INFO")
                add_log(f"Merged DataFrame has {len(merged_df)} rows", "DEBUG")
                add_log(f"Output directory: {output_dir}", "DEBUG")
                
                # Rename speakers to match your format
                merged_df_for_chunking = merged_df.copy()
                add_log("Renaming speakers from AssemblyAI format...", "DEBUG")
                merged_df_for_chunking['Speaker'] = merged_df_for_chunking['Speaker'].apply(
                    lambda x: "Mysteryshopper" if x in ['SPEAKER_A', 'SPEAKER_0'] else "InsuranceAgent"
                )
                
                # Create chunks
                add_log("Calling create_conversation_chunks()...", "DEBUG")
                chunked_df = create_conversation_chunks(merged_df_for_chunking)
                
                if chunked_df is not None and len(chunked_df) > 0:
                    add_log(f"Created {len(chunked_df)} chunks successfully", "INFO")
                    
                    # Save only Excel file with quality analysis
                    chunked_excel = Path(output_dir) / f"{base}_full_chunked_with_quality.xlsx"
                    
                    add_log(f"Creating chunked Excel with quality analysis: {chunked_excel}", "DEBUG")
                    
                    # Run quality analysis on chunked data
                    chunked_quality_results = analyze_chunked_quality(chunked_df)
                    
                    # Create Excel with quality analysis
                    create_chunked_excel_with_quality(chunked_df, chunked_quality_results, chunked_excel, audio_file, len(merged_df))
                    
                    add_log(f"Created chunked Excel with quality analysis: {chunked_excel}")
                    
                else:
                    add_log("Chunking function returned None or empty DataFrame", "ERROR")
                    
            except Exception as e:
                add_log(f"CHUNKING ERROR: {str(e)}", "ERROR")
                add_log(f"Chunking traceback: {traceback.format_exc()}", "ERROR")
            
        return True, log
        
    except Exception as e:
        import traceback
        error_details = f"Exception during transcription: {str(e)}\nTraceback: {traceback.format_exc()}"
        add_log(error_details, "ERROR")
        return False, error_details

def create_transcript_with_quality_excel(merged_csv_path, quality_results, label_col, output_folder, base_name):
    """Create Excel file with merged transcript and quality analysis logs as separate sheets."""
    import pandas as pd
    import os
    excel_path = os.path.join(output_folder, f"{base_name}_with_quality.xlsx")
    merged_df = pd.read_csv(merged_csv_path)
    
    # Prepare DataFrames for logs
    flagged_df = pd.DataFrame(quality_results['flagged_segments'])
    long_df = pd.DataFrame(quality_results['long_segments'])
    gap_df = pd.DataFrame(quality_results['gap_segments'])

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        # Sheet 1: Full transcript
        merged_df.to_excel(writer, sheet_name='Transcript', index=False)
        # Sheet 2: Quality Analysis Logs
        startrow = 0
        # Flagged segments
        if not flagged_df.empty:
            flagged_df[["Start Time", "End Time", "Speaker", "Text", label_col]].to_excel(
                writer, sheet_name='Quality_Logs', index=False, startrow=startrow
            )
            startrow += len(flagged_df) + 2  # +2 for header and one blank row
        # Long segments
        if not long_df.empty:
            long_df[["Start Time", "End Time", "Speaker", "Text", label_col]].to_excel(
                writer, sheet_name='Quality_Logs', index=False, startrow=startrow
            )
            startrow += len(long_df) + 2
        # Gap segments
        if not gap_df.empty:
            gap_df[["Start Time", "End Time", "Speaker", "Text", label_col]].to_excel(
                writer, sheet_name='Quality_Logs', index=False, startrow=startrow
            )
    return excel_path

def process_large_audio_file(audio_file_path, output_folder, st, uploaded_file, progress_bar, status_text):
    """Process a large audio file with progress tracking"""
    try:
        base_name = os.path.splitext(uploaded_file.name)[0]
        
        # Store base_name in session state
        st.session_state['base_name'] = base_name
        
        status_text.text("Running transcription...")
        progress_bar.progress(0.1)

        # Run transcription using AssemblyAI with quality analysis
        success_tx, tx_logs = run_elevenlabs_transcription(audio_file_path, output_folder)
        
        # Store transcription logs
        st.session_state.tx_logs = tx_logs
        add_log("Transcription completed", "INFO")
        
        if success_tx:
            progress_bar.progress(0.5)
            status_text.text("Transcription finished – preparing output...")
            
            # Find the merged CSV file (AssemblyAI creates _merged.csv, not _full_merged.csv)
            merged_csv_path = os.path.join(output_folder, f"{base_name}_merged.csv")
            
            if os.path.exists(merged_csv_path):
                add_log(f"Found merged CSV: {merged_csv_path}")
                
                # Quality analysis was already done in run_elevenlabs_transcription
                # Just read the results for download preparation
                try:
                    df = pd.read_csv(merged_csv_path)
                    label_col = 'Label' if 'Label' in df.columns else ('Notes' if 'Notes' in df.columns else None)
                    
                    # Get quality results from session state (set by run_elevenlabs_transcription)
                    quality_results = st.session_state.get('quality_results', {
                        'flagged_segments': [],
                        'long_segments': [],
                        'gap_segments': [],
                        'total_issues': 0,
                        'total_segments': len(df)
                    })
                    
                except Exception as e:
                    add_log(f"Error reading quality results: {str(e)}", "ERROR")
                    quality_results = {'flagged_segments': [], 'long_segments': [], 'gap_segments': [], 'total_issues': 0, 'total_segments': 0}
                    label_col = None
                
                # Create output files dictionary
                output_files = {'merged_csv': merged_csv_path}
                
                # Only use chunked files with quality analysis
                chunked_excel_quality = os.path.join(output_folder, f"{base_name}_chunked_with_quality.xlsx")
                chunked_csv_quality = os.path.join(output_folder, f"{base_name}_chunked.csv")
                chunked_txt_quality = os.path.join(output_folder, f"{base_name}_chunked.txt")
                
                # Only include quality analysis versions
                if os.path.exists(chunked_excel_quality):
                    output_files['chunked_excel'] = chunked_excel_quality
                    add_log(f"Found chunked Excel with quality analysis: {chunked_excel_quality}")
                
                if os.path.exists(chunked_csv_quality):
                    output_files['chunked_csv'] = chunked_csv_quality
                    add_log(f"Found chunked CSV with quality analysis: {chunked_csv_quality}")
                    
                if os.path.exists(chunked_txt_quality):
                    output_files['chunked_txt'] = chunked_txt_quality
                    add_log(f"Found chunked TXT: {chunked_txt_quality}")
                
                excel_with_quality = create_transcript_with_quality_excel(merged_csv_path, quality_results, label_col, output_folder, base_name)
                output_files['excel_with_quality'] = excel_with_quality

                st.success("Processing completed successfully!")

                # Create download with only the main chunked Excel file with quality analysis
                if 'chunked_excel' in output_files:
                    # Direct download of single Excel file (no ZIP needed)
                    with open(output_files['chunked_excel'], 'rb') as f:
                        excel_data = f.read()
                    
                    st.session_state['download_data'] = excel_data
                    st.session_state['download_filename'] = os.path.basename(output_files['chunked_excel'])
                    st.session_state['download_mime'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    add_log(f"Prepared single file download: {os.path.basename(output_files['chunked_excel'])}")
                else:
                    add_log("CHUNKING COMPLETE! Excel file with quality analysis created", "INFO")
                    st.session_state['download_data'] = None
                    
                # Store results in session state to prevent UI disappearing
                st.session_state['processing_complete'] = True
                st.session_state['quality_results'] = quality_results
                st.session_state['output_files'] = output_files
                st.session_state['show_results'] = True
                st.session_state['base_name'] = base_name

                # Show preview - prioritize chunked format if available
                if 'chunked_excel' in output_files:
                    st.subheader("Chunked Transcript Preview (5 segments per chunk)")
                    try:
                        chunked_preview = pd.read_excel(output_files['chunked_excel'], sheet_name='Chunked_Transcript').head(5)
                        st.dataframe(chunked_preview, use_container_width=True)
                    except Exception as e:
                        st.warning(f"Could not load chunked preview: {str(e)}")
                        # Try to load without sheet name
                        try:
                            chunked_preview = pd.read_excel(output_files['chunked_excel']).head(5)
                            st.dataframe(chunked_preview, use_container_width=True)
                        except Exception as e2:
                            st.error(f"Could not load Excel file: {str(e2)}")
                    
                    # Show sample of combined text
                    if len(chunked_preview) > 0:
                        st.subheader("Sample Chunk with Speaker Labels")
                        sample_text = chunked_preview.iloc[0]['Combined_Text']
                        st.text_area("First chunk content:", sample_text, height=150)
                        
                    # Show quality summary for chunks
                    try:
                        quality_summary = pd.read_excel(output_files['chunked_excel'], sheet_name='Quality_Summary')
                        st.subheader("Chunked Quality Analysis")
                        st.dataframe(quality_summary, use_container_width=True)
                    except:
                        pass
                elif 'merged_csv' in output_files:
                    st.subheader("Merged Transcript Preview (Individual Segments)")
                    df_preview = pd.read_csv(output_files['merged_csv']).head(10)
                    st.dataframe(df_preview, use_container_width=True)
                    st.warning("Chunked files not found - showing merged transcript instead")
                
                progress_bar.progress(1.0)
                status_text.text("Processing complete!")
                add_log("Processing completed successfully!")
                
            else:
                error_msg = "Merged CSV file not found after transcription"
                st.error(f"{error_msg}")
                add_log(error_msg, "ERROR")
        else:
            error_msg = f"Transcription failed: {tx_logs}"
            st.error(f"{error_msg}")
            add_log(error_msg, "ERROR")
            # Display detailed error information
            if tx_logs and isinstance(tx_logs, str):
                st.error(f"Error details: {tx_logs}")
                with st.expander("Full Error Details", expanded=True):
                    st.text(tx_logs)
            
    except Exception as e:
        error_msg = f"Processing failed: {str(e)}"
        st.error(f"{error_msg}")
        add_log(error_msg, "ERROR")
        add_log(traceback.format_exc(), "ERROR")
        raise

# -------------------------------------------------------------
# Helper: flag long gaps with no transcription
# -------------------------------------------------------------

def flag_long_gaps(csv_path, gap_seconds: int = 60):
    """Add a warning to Label/Notes if there is a silent gap longer than gap_seconds
    between consecutive segments.
    """
    try:
        import pandas as pd
        import re
        
        df = pd.read_csv(csv_path)
        label_col = 'Label' if 'Label' in df.columns else ('Notes' if 'Notes' in df.columns else None)
        if label_col is None:
            df['Label'] = ''
            label_col = 'Label'

        # Parse timestamps 00:01:20,519 -> timedelta
        def to_timedelta(t):
            try:
                t = t.strip()
                # allow comma or . between seconds and ms and optional spaces after comma
                m = re.match(r"(?P<h>\d{1,2}):(?P<m>\d{1,2}):(?P<s>\d{1,2})[,.]\s*(?P<ms>\d{1,3})", t)
                if not m:
                    return float('nan')
                total = int(m['h'])*3600 + int(m['m'])*60 + int(m['s']) + int(m['ms'])/1000
                return total
            except Exception:
                return None

        # Helper to append messages safely
        def add_flag(series, mask, flag_text):
            def _append(val):
                base = '' if (pd.isna(val) or str(val).strip()=='') else f"{str(val).strip()}; "
                return base + flag_text
            # Use .loc with proper assignment to avoid warning
            series_copy = series.copy()
            series_copy.loc[mask] = series_copy.loc[mask].apply(_append)
            return series_copy

        starts = df['Start Time'].apply(lambda x: to_timedelta(str(x)))
        prev_ends = df['End Time'].shift().apply(lambda x: to_timedelta(str(x)))
        gaps = starts - prev_ends
        df['__gap__'] = gaps
        # First row will be NaN; ignore
        gap_mask = df['__gap__'] > gap_seconds
        if gap_mask.any():
            df[label_col] = add_flag(df[label_col], gap_mask, '>60s silence, missing audio segment')
        # After gap detection modifications within flag_long_gaps
        # Extra: mark empty or whitespace-only text
        empty_mask = df['Text'].apply(lambda x: str(x).strip() == '')
        if empty_mask.any():
            df[label_col] = add_flag(df[label_col], empty_mask, 'empty text, missing transcription')
        # Flag rows whose own duration > 120 s (possible untranscribed part)
        durations = df.apply(lambda r: to_timedelta(str(r['End Time'])) - to_timedelta(str(r['Start Time'])), axis=1)
        longgap_mask = durations > 120
        if longgap_mask.any():
            df[label_col] = add_flag(df[label_col], longgap_mask, '>120s single segment, possible missing transcription')
        df.drop(columns='__gap__', inplace=True)
        df.to_csv(csv_path, index=False)
        return gap_mask.sum()
    except Exception as e:
        print(f"Gap validation failed for {csv_path}: {e}")
        return 0

# -------------------------------------------------------------
# Helper: detect and flag 90+ second gaps in chunked CSV
# -------------------------------------------------------------

def detect_large_gaps_in_transcript(csv_path, min_gap_seconds=90):
    """Detect gaps longer than min_gap_seconds in the transcript that may indicate missing transcription"""
    try:
        df = pd.read_csv(csv_path)
        
        def to_seconds(timestamp_str):
            """Convert HH:MM:SS,mmm to seconds"""
            try:
                time_part, ms_part = timestamp_str.split(',')
                h, m, s = map(int, time_part.split(':'))
                ms = int(ms_part)
                return h * 3600 + m * 60 + s + ms / 1000
            except:
                return 0
        
        large_gaps = []
        
        for i in range(len(df) - 1):
            current_end = to_seconds(df.iloc[i]['End Time'])
            next_start = to_seconds(df.iloc[i + 1]['Start Time'])
            gap_duration = next_start - current_end
            
            if gap_duration >= min_gap_seconds:
                large_gaps.append({
                    'Gap Number': len(large_gaps) + 1,
                    'After Segment': i + 1,
                    'Gap Start': df.iloc[i]['End Time'],
                    'Gap End': df.iloc[i + 1]['Start Time'],
                    'Duration': format_duration(gap_duration),
                    'duration_seconds': gap_duration,
                    'Previous Text': df.iloc[i]['Text'][:50] + "..." if len(df.iloc[i]['Text']) > 50 else df.iloc[i]['Text'],
                    'Next Text': df.iloc[i + 1]['Text'][:50] + "..." if len(df.iloc[i + 1]['Text']) > 50 else df.iloc[i + 1]['Text']
                })
        
        return {
            'gaps_found': len(large_gaps) > 0,
            'large_gaps': large_gaps,
            'total_gaps': len(large_gaps)
        }
    except Exception as e:
        return {
            'gaps_found': False,
            'large_gaps': [],
            'total_gaps': 0,
            'error': str(e)
        }

def format_duration(seconds):
    """Format seconds into readable duration (e.g., '2m 30s' or '1h 15m 30s')"""
    if seconds < 60:
        return f"{int(seconds)}s"
    elif seconds < 3600:
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes}m {secs}s" if secs > 0 else f"{minutes}m"
    else:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        if secs > 0:
            return f"{hours}h {minutes}m {secs}s"
        elif minutes > 0:
            return f"{hours}h {minutes}m"
        else:
            return f"{hours}h"

def detect_90s_gaps_in_chunked_csv(csv_path):
    """
    Detect chunks longer than 90 seconds in the chunked CSV file, which may indicate
    missing transcription within the chunk. Print timestamps and flag them in Notes.
    Returns count of long chunks found and log messages.
    """
    log_messages = []
    try:
        df = pd.read_csv(csv_path)
        
        # Parse timestamps to seconds
        def to_seconds(timestamp_str):
            try:
                timestamp_str = str(timestamp_str).strip().strip('"')
                import re
                # Match HH:MM:SS,mmm or HH:MM:SS.mmm format
                match = re.match(r'(\d{1,2}):(\d{1,2}):(\d{1,2})[,.](\d{1,3})', timestamp_str)
                if match:
                    hours, minutes, seconds, milliseconds = match.groups()
                    total_seconds = int(hours) * 3600 + int(minutes) * 60 + int(seconds) + int(milliseconds) / 1000
                    return total_seconds
                return None
            except:
                return None
        
        # Convert timestamps to seconds
        df['start_seconds'] = df['Start Time'].apply(to_seconds)
        df['end_seconds'] = df['End Time'].apply(to_seconds)
        
        # Calculate duration of each chunk
        long_chunks_found = []
        long_chunk_count = 0
        
        for i in range(len(df)):
            start_sec = df.iloc[i]['start_seconds']
            end_sec = df.iloc[i]['end_seconds']
            
            if start_sec is not None and end_sec is not None:
                chunk_duration = end_sec - start_sec
                
                if chunk_duration >= 90:  # 90+ second chunk
                    long_chunk_count += 1
                    start_time = df.iloc[i]['Start Time']
                    end_time = df.iloc[i]['End Time']
                    
                    # Format chunk duration
                    duration_minutes = int(chunk_duration // 60)
                    duration_remainder = int(chunk_duration % 60)
                    duration_str = f"{duration_minutes}m {duration_remainder}s" if duration_minutes > 0 else f"{duration_remainder}s"
                    
                    chunk_info = {
                        'chunk_number': long_chunk_count,
                        'start_time': start_time,
                        'end_time': end_time,
                        'duration': duration_str,
                        'duration_seconds': chunk_duration
                    }
                    long_chunks_found.append(chunk_info)
                    
                    # Print and log the chunk information
                    log_msg = f"Long Chunk #{long_chunk_count}: {duration_str} chunk from {start_time} to {end_time}"
                    print(log_msg)
                    log_messages.append(log_msg)
                    
                    # Flag in Notes column
                    current_notes = df.iloc[i]['Notes']
                    if pd.isna(current_notes) or str(current_notes).strip() == '' or current_notes == 'nan':
                        flag_text = f"Long chunk ({duration_str}) - may contain missing transcription"
                    else:
                        flag_text = f"{str(current_notes).strip()}; Long chunk ({duration_str}) - may contain missing transcription"
                    
                    df.iloc[i, df.columns.get_loc('Notes')] = flag_text
        
        # Save the updated CSV with flags
        df.drop(columns=['start_seconds', 'end_seconds'], inplace=True)
        df.to_csv(csv_path, index=False)
        
        # Print and log summary
        if long_chunk_count > 0:
            summary_msg = f"Summary: Found {long_chunk_count} chunk(s) longer than 90 seconds"
            print(summary_msg)
            log_messages.append(summary_msg)
            for chunk in long_chunks_found:
                detail_msg = f"   • Chunk {chunk['chunk_number']}: {chunk['duration']} from {chunk['start_time']} to {chunk['end_time']}"
                print(detail_msg)
                log_messages.append(detail_msg)
        else:
            success_msg = "No chunks longer than 90 seconds found"
            print(success_msg)
            log_messages.append(success_msg)
        
        return long_chunk_count, long_chunks_found, log_messages
        
    except Exception as e:
        error_msg = f"Error detecting long chunks in chunked CSV {csv_path}: {e}"
        print(error_msg)
        log_messages.append(error_msg)
        return 0, [], log_messages

# -------------------------------------------------------------

def display_quality_analysis():
    """Display the persistent quality analysis results"""
    if hasattr(st.session_state, 'quality_results') and st.session_state.quality_results:
        st.subheader("Transcript Quality Analysis")
        
        quality_results = st.session_state.quality_results
        label_col = st.session_state.get('label_col', 'Notes')
        
        # Display flagged segments
        if quality_results['flagged_segments']:
            st.warning(f"{len(quality_results['flagged_segments'])} segment(s) were flagged as requiring human transcription due to gibberish or unsupported language.")
            with st.expander("View flagged segments"):
                df_flagged = pd.DataFrame(quality_results['flagged_segments'])
                st.dataframe(df_flagged[['Start Time', 'End Time', 'Speaker', 'Text', label_col]], use_container_width=True)
        
        # Display long segments
        if quality_results['long_segments']:
            st.error(f"Found {len(quality_results['long_segments'])} segment(s) longer than 120 seconds - possible missing transcription!")
            with st.expander("View long segments"):
                df_long = pd.DataFrame(quality_results['long_segments'])
                st.dataframe(df_long[['Start Time', 'End Time', 'Speaker', 'Text', label_col]], use_container_width=True)
        
        # Display gap segments
        if quality_results['gap_segments']:
            st.warning(f"Found {len(quality_results['gap_segments'])} segment(s) with gaps > 60s")
            with st.expander("View segments with gaps"):
                df_gaps = pd.DataFrame(quality_results['gap_segments'])
                st.dataframe(df_gaps[['Start Time', 'End Time', 'Speaker', 'Text', label_col]], use_container_width=True)
        
        # Display summary
        if quality_results['total_issues'] == 0:
            st.success("No quality issues detected!")
        else:
            st.info(f"Quality Summary: {quality_results['total_issues']} total issues detected across {quality_results['total_segments']} segments")

def main():
    st.set_page_config(page_title="Intage Audio Transcription UI", layout="wide")
    
    # Initialize logging
    init_logging()
    
    st.title("Intage Audio Transcription UI")
    
    # File uploader section
    st.subheader("Choose an audio file")
    uploaded_file = st.file_uploader(
        "Drag and drop file here",
        type=["wav", "mp3", "flac", "ogg", "m4a"],
        help="Limit 500MB per file • WAV, MP3, FLAC, OGG, M4A"
    )

    if uploaded_file:
        file_size_mb = uploaded_file.size / (1024 * 1024)
        st.success(f"File uploaded: {uploaded_file.name} ({file_size_mb:.1f} MB)")
        
        # Create temp directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save uploaded file
            temp_audio_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_audio_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            # Create output directory
            output_dir = os.path.join(temp_dir, "output")
            os.makedirs(output_dir, exist_ok=True)
            
            # Processing method selection
            st.subheader("Processing method:")
            transcription_mode = st.radio(
                "Select transcription mode",
                ["Transcribe whole file"],
                label_visibility="collapsed"
            )
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            if st.button("Start Transcription", type="primary", key="start_transcription"):
                try:
                    process_large_audio_file(temp_audio_path, output_dir, st, uploaded_file, progress_bar, status_text)
                except Exception as e:
                    st.error(f"Processing failed: {str(e)}")
                    st.error(f"Traceback: {traceback.format_exc()}")
                    add_log(f"Traceback: {traceback.format_exc()}", "ERROR")
            
    # Display persistent results outside temp directory context
    # This prevents UI from disappearing after download
    if hasattr(st.session_state, 'zip_data') and st.session_state.zip_data:
        st.subheader("Processing Results")
        
        # Show download button with persistent data
        base_name = st.session_state.get('base_name', 'transcript')
        download_key = f"persistent_download_{base_name}"
        st.download_button(
            label="Download Transcript + Quality Logs (Excel in ZIP)",
            data=st.session_state.zip_data,
            file_name=f"{base_name}_transcript_quality.zip",
            mime="application/zip",
            type="primary",
            key=download_key
        )
        
        # Show preview if available
        if hasattr(st.session_state, 'output_files') and st.session_state.output_files:
            output_files = st.session_state.output_files
            
            if 'chunked_excel' in output_files and os.path.exists(output_files['chunked_excel']):
                st.subheader("✅ Chunked Transcript Preview (5 segments per chunk)")
                try:
                    chunked_preview = pd.read_excel(output_files['chunked_excel'], sheet_name='Chunked_Transcript').head(5)
                    st.dataframe(chunked_preview, use_container_width=True)
                    
                    # Show sample of combined text and Notes column
                    if len(chunked_preview) > 0:
                        st.subheader("Sample Chunk with Speaker Labels")
                        sample_text = chunked_preview.iloc[0]['Combined_Text']
                        st.text_area("First chunk content:", value=sample_text, height=150, disabled=True)
                        
                        # Show Notes column if it exists
                        if 'Notes' in chunked_preview.columns:
                            st.subheader("Quality Analysis Notes")
                            notes_preview = chunked_preview[['Chunk_Number', 'Notes']].dropna(subset=['Notes'])
                            if not notes_preview.empty:
                                st.dataframe(notes_preview, use_container_width=True)
                            else:
                                st.info("No quality issues detected in preview chunks")
                        
                    # Show quality summary for chunks
                    try:
                        quality_summary = pd.read_excel(output_files['chunked_excel'], sheet_name='Quality_Summary')
                        st.subheader("Chunked Quality Analysis")
                        st.dataframe(quality_summary, use_container_width=True)
                    except:
                        pass
                except Exception as e:
                    st.warning(f"Could not load chunked preview: {str(e)}")
                    add_log(f"Preview error: {str(e)}", "ERROR")
                    
            elif 'merged_csv' in output_files and os.path.exists(output_files['merged_csv']):
                st.subheader("Merged Transcript Preview (Individual Segments)")
                try:
                    df_preview = pd.read_csv(output_files['merged_csv']).head(10)
                    st.dataframe(df_preview, use_container_width=True)
                    st.warning("Chunked files not found - showing merged transcript instead")
                except Exception as e:
                    st.warning(f"Could not load merged preview: {str(e)}")
    
    # Show persistent download button and results after processing
    if st.session_state.get('show_results', False) and st.session_state.get('download_data'):
        st.markdown("---")
        st.subheader("Processing Results")
        
        # Download button that persists - single Excel file
        if st.session_state.get('download_data'):
            download_key = f"persistent_download_{st.session_state.get('base_name', 'transcript')}"
            st.download_button(
                label="Download Chunked Transcript + Quality Analysis (Excel)",
                data=st.session_state['download_data'],
                file_name=st.session_state.get('download_filename', 'chunked_with_quality.xlsx'),
                mime=st.session_state.get('download_mime', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'),
                type="primary",
                key=download_key
            )
        
        # Show file contents info
        if st.session_state.get('output_files') and 'chunked_excel' in st.session_state['output_files']:
            st.info("**Download includes:**")
            st.write("• Chunked Excel with quality analysis sheets")
            st.write("• Notes column with quality flags (gaps >240 seconds)")
            st.write("• Quality summary and issue details")
    
    # Display quality analysis if it exists (persists after download)
    if hasattr(st.session_state, 'quality_results') and st.session_state.quality_results:
        display_quality_analysis()
    
    # Clear session button
    if st.button("Clear session & temp files"):
        # Clear all session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.experimental_rerun()

if __name__ == "__main__":
    main() 